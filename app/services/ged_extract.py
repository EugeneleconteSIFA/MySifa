"""MySifa - Extraction de texte pour l'indexation plein texte de la GED qualite.

Utilise uniquement des bibliotheques deja presentes dans requirements.txt :
pdfplumber, python-docx, openpyxl. Aucune dependance nouvelle.

Contrat : extract_text(path, ext) -> (texte, statut)
  statut = 'ok'      : du texte a ete extrait
           'skipped' : type non supporte, ou fichier sans couche texte
                       (typiquement un PDF scanne -> il faudra le taguer a la main,
                        ou brancher un OCR plus tard sans changer de schema)
           'error'   : le parseur a leve une exception

Garde-fous : on plafonne volontairement le volume extrait. Un rapport d'audit de
300 pages n'a pas besoin d'etre indexe integralement pour etre retrouve, et un
index qui explose ralentit toutes les recherches.
"""
from __future__ import annotations

import re

MAX_CHARS = 500_000      # au-dela, on tronque : ca suffit tres largement
MAX_PDF_PAGES = 60
MAX_XLSX_SHEETS = 20
MAX_XLSX_ROWS = 5_000

# Extensions dont on sait extraire du texte
TEXT_EXTS = {"txt", "md", "csv", "tsv", "log", "json", "xml", "html", "htm"}
SUPPORTED = TEXT_EXTS | {"pdf", "docx", "xlsx", "xlsm", "pptx"}


def _clean(txt: str) -> str:
    """Normalise les blancs : l'index n'a que faire des tabulations et des
    triples sauts de ligne, et ca divise la taille stockee par deux."""
    if not txt:
        return ""
    txt = txt.replace("\x00", " ")
    txt = re.sub(r"[ \t ]+", " ", txt)
    txt = re.sub(r"\n{3,}", "\n\n", txt)
    return txt.strip()[:MAX_CHARS]


def _from_pdf(path: str) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            if i >= MAX_PDF_PAGES:
                break
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
            if sum(len(p) for p in parts) > MAX_CHARS:
                break
    return "\n".join(parts)


def _from_docx(path: str) -> str:
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    # Les tableaux portent souvent l'essentiel dans les documents qualite
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_xlsx(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    parts = []
    try:
        for si, ws in enumerate(wb.worksheets):
            if si >= MAX_XLSX_SHEETS:
                break
            parts.append(f"[{ws.title}]")
            for ri, row in enumerate(ws.iter_rows(values_only=True)):
                if ri >= MAX_XLSX_ROWS:
                    break
                vals = [str(v) for v in row if v is not None and str(v).strip()]
                if vals:
                    parts.append(" | ".join(vals))
                if sum(len(p) for p in parts) > MAX_CHARS:
                    break
            if sum(len(p) for p in parts) > MAX_CHARS:
                break
    finally:
        try:
            wb.close()
        except Exception:
            pass
    return "\n".join(parts)


def _from_pptx(path: str) -> str:
    """python-pptx n'est pas dans requirements.txt : on lit le XML brut du zip.
    Suffisant pour l'indexation, et zero dependance nouvelle."""
    import zipfile
    parts = []
    with zipfile.ZipFile(path) as z:
        names = [n for n in z.namelist()
                 if n.startswith("ppt/slides/slide") and n.endswith(".xml")]
        for n in sorted(names):
            try:
                xml = z.read(n).decode("utf-8", "ignore")
            except Exception:
                continue
            parts.extend(re.findall(r"<a:t>(.*?)</a:t>", xml, re.S))
    return "\n".join(parts)


def _from_text(path: str) -> str:
    with open(path, "rb") as fh:
        raw = fh.read(MAX_CHARS * 2)
    for enc in ("utf-8", "cp1252", "latin-1"):
        try:
            return raw.decode(enc)
        except Exception:
            continue
    return raw.decode("utf-8", "ignore")


def extract_text(path: str, ext: str) -> tuple[str, str]:
    """Renvoie (texte_extrait, statut)."""
    ext = (ext or "").lower().lstrip(".")
    if ext not in SUPPORTED:
        return "", "skipped"
    try:
        if ext == "pdf":
            txt = _from_pdf(path)
        elif ext == "docx":
            txt = _from_docx(path)
        elif ext in ("xlsx", "xlsm"):
            txt = _from_xlsx(path)
        elif ext == "pptx":
            txt = _from_pptx(path)
        else:
            txt = _from_text(path)
    except Exception:
        return "", "error"

    txt = _clean(txt)
    # Un PDF scanne renvoie une chaine vide ou quelques caracteres parasites :
    # on le marque 'skipped' pour que l'UI puisse dire "contenu non indexe".
    if len(txt) < 12:
        return "", "skipped"
    return txt, "ok"
