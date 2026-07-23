# -*- coding: utf-8 -*-
"""
SIFA — Génération DOCX des documents officiels (Déclarations UE de Conformité,
etc.). Miroir de app/services/sifa_doc_pdf.py, utilisant python-docx.

Ce service produit un fichier Word éditable à partir des mêmes sections
(SECTIONS_META) et des mêmes textes par défaut (SEC_*_BODY) que la version
PDF. La priorité de résolution des textes est identique :

  1. version.sections_overrides[sec_id].custom_body   (par client)
  2. template.default_body_overrides[sec_id]          (par template — admin)
  3. SEC_*_BODY hardcodé dans sifa_doc_pdf             (défaut usine)

Utilisé par app/routers/qualite.py — route
GET /api/qualite/sifa-docs/versions/{vid}/docx.
"""

from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.services.sifa_doc_pdf import (
    SECTIONS_META,
    SEC_2_BODY, SEC_4_INTRO,
    SEC_4_1_BODY, SEC_4_2_BODY, SEC_4_3_BODY, SEC_4_4_BODY, SEC_4_4_NOTE,
    SEC_4_5_BODY, SEC_4_6_BODY, SEC_4_7_BODY, SEC_4_8_BODY, SEC_4_9_BODY,
    SEC_5_BODY, SEC_6_BODY, SEC_7_BODY, SEC_8_BODY,
    _locate_logo, _locate_cachet, _locate_signature,
)

# ─── Palette (miroir du PDF) ─────────────────────────────────────────────
NAVY = RGBColor(0x0F, 0x17, 0x2A)
GREY = RGBColor(0x64, 0x74, 0x8B)
MUTED = RGBColor(0x94, 0xA3, 0xB8)
BORDER_HEX = "CBD5E1"
LIGHT_BG_HEX = "F8FAFC"
YELLOW_HEX = "FBBF24"
YELLOW_BG_HEX = "FEF9E7"

# ─── Mois en français ─────────────────────────────────────────────────────
_MONTHS_FR = [
    "Janvier", "Février", "Mars", "Avril", "Mai", "Juin",
    "Juillet", "Août", "Septembre", "Octobre", "Novembre", "Décembre",
]


def _fr_month_year(iso_str: str) -> str:
    try:
        d = datetime.fromisoformat(iso_str)
        return f"{_MONTHS_FR[d.month - 1]} {d.year}"
    except Exception:
        now = datetime.now()
        return f"{_MONTHS_FR[now.month - 1]} {now.year}"


def _fr_date_parts(iso_str: str):
    try:
        d = datetime.fromisoformat(iso_str)
        return f"{d.day:02d}", f"{d.month:02d}", str(d.year)
    except Exception:
        return "____", "____", str(datetime.now().year)


# ─── Helpers de bas niveau python-docx ────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    """Applique une couleur de fond à une cellule de tableau."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_border(cell, color_hex: str = BORDER_HEX, size: str = "6"):
    """Applique une bordure fine grise à une cellule."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _add_para(doc_or_cell, text: str, *,
              bold: bool = False, italic: bool = False,
              size: float = 10.0, color: RGBColor = NAVY,
              align=None, space_after: float = 4.0, space_before: float = 0.0):
    """Ajoute un paragraphe avec un run principal."""
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    run = p.add_run(text or "")
    run.bold = bold
    run.italic = italic
    font = run.font
    font.name = "Calibri"
    font.size = Pt(size)
    font.color.rgb = color
    return p


def _add_heading_h2(doc, text: str, level_size: float = 13.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text or "")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(level_size)
    run.font.color.rgb = NAVY
    # Ligne fine sous le titre
    _add_bottom_border(p)
    return p


def _add_heading_h3(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text or "")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    return p


def _add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BORDER_HEX)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_body(doc_or_cell, text: str, *, italic: bool = False,
              size: float = 10.0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    return _add_para(doc_or_cell, text, italic=italic, size=size,
                     color=NAVY, align=align, space_after=4.0)


def _add_bullet(doc, text: str):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("• " + (text or ""))
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    return p


def _kv_table(doc, rows, col_widths_cm=(4.5, 12.5)):
    """Tableau clé/valeur (1ère colonne fond clair, bordures grises)."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.autofit = False
    for i, (k, v) in enumerate(rows):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(col_widths_cm[0])
        c1.width = Cm(col_widths_cm[1])
        _set_cell_bg(c0, LIGHT_BG_HEX)
        _set_cell_border(c0)
        _set_cell_border(c1)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Clear default paragraph
        c0.text = ""
        c1.text = ""
        pk = c0.paragraphs[0]
        rk = pk.add_run(k or "")
        rk.bold = True
        rk.font.name = "Calibri"
        rk.font.size = Pt(10)
        rk.font.color.rgb = NAVY
        pv = c1.paragraphs[0]
        rv = pv.add_run(v or "")
        rv.font.name = "Calibri"
        rv.font.size = Pt(10)
        rv.font.color.rgb = NAVY
    return tbl


# ─── Résolution du body avec priorité ─────────────────────────────────────
def _resolve_body(sec_id: str, default_body: str,
                  version_overrides: dict,
                  template_overrides: dict) -> str:
    """Applique la priorité : version.custom_body > template.default_override
    > SEC_*_BODY hardcodé."""
    v = (version_overrides or {}).get(sec_id, {})
    v_custom = v.get("custom_body") if isinstance(v, dict) else None
    if v_custom and str(v_custom).strip():
        return str(v_custom).strip()
    t_default = (template_overrides or {}).get(sec_id)
    if t_default and str(t_default).strip():
        return str(t_default).strip()
    return default_body or ""


def _section_included(sec, version_overrides: dict) -> bool:
    if not sec.get("removable"):
        return True
    ov = (version_overrides or {}).get(sec["id"], {})
    if not isinstance(ov, dict):
        return True
    return ov.get("include") is not False


# ─── Header / Footer via section.header/footer ────────────────────────────
def _install_header_footer(doc, ref: str, header_date: str):
    section = doc.sections[0]
    section.top_margin = Mm(30)
    section.bottom_margin = Mm(30)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    # Header : logo + libellé document + réf/date
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_path = _locate_logo()
    if logo_path:
        try:
            run_logo = header_para.add_run()
            run_logo.add_picture(logo_path, width=Mm(28))
            header_para.add_run("\t\t")
        except Exception:
            pass
    r1 = header_para.add_run("SIFA — Déclaration UE de Conformité")
    r1.italic = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = GREY

    header_para2 = header.add_paragraph()
    header_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = header_para2.add_run(f"{header_date}   Réf. {ref}")
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = GREY

    # Footer : coordonnées SIFA + page
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("SIFA · 45 rue Rollin · 59100 Roubaix · France")
    fr.font.size = Pt(9)
    fr.font.color.rgb = GREY
    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr2 = fp2.add_run("+33 (0)3 20 69 01 01 · commandes@sifa.pro")
    fr2.font.size = Pt(9)
    fr2.font.color.rgb = GREY


# ─── Blocs par section ────────────────────────────────────────────────────
def _title_block_docx(doc, ref: str, date_emission_iso: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DÉCLARATION UE DE CONFORMITÉ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("EU Declaration of Conformity (DoC)")
    r2.italic = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = GREY
    p2.paragraph_format.space_after = Pt(14)

    jour, mois, annee = _fr_date_parts(date_emission_iso)
    _kv_table(doc, [
        ("Numéro d'identification unique", ref or ""),
        ("Date d'émission", f"Roubaix, le {jour} / {mois} / {annee}"),
    ])
    doc.add_paragraph()  # espacement


def _sec_1_docx(doc, num=1):
    _add_heading_h2(doc, f"{num}. Fabricant")
    _kv_table(doc, [
        ("Raison sociale", "SIFA"),
        ("Adresse", "45 rue Rollin, 59100 Roubaix, France"),
        ("Activité", "Fabricant d'étiquettes adhésives"),
    ])


def _sec_2_docx(doc, body, num=2):
    _add_heading_h2(doc, f"{num}. Nature de l'activité SIFA")
    _add_body(doc, body or SEC_2_BODY)


def _sec_3_docx(doc, ctx, num=3):
    _add_heading_h2(doc, f"{num}. Fournisseurs — origine géographique")
    client_nom = ctx["client_nom"]
    fournisseurs = ctx["fournisseurs"] or []
    if client_nom:
        _add_body(doc, f"Pour les étiquettes livrées à {client_nom}, les "
                       f"matières entrant dans les étiquettes proviennent des "
                       f"fournisseurs suivants :")
    else:
        _add_body(doc, "Les matières entrant dans les étiquettes proviennent "
                       "des fournisseurs suivants :")
    for f in fournisseurs:
        nom = (f.get("nom") or "").strip()
        pays = (f.get("pays_origine") or "").strip() or "origine à préciser"
        _add_bullet(doc, f"{nom} — matières fabriquées en {pays}.")
    _add_body(doc,
              "Toutes les matières à l'origine de la fabrication des "
              "étiquettes couvertes par la présente Déclaration sont "
              "fabriquées au sein de l'Union européenne ou du Royaume-Uni.")


def _sec_4_intro_docx(doc, body, num=4):
    _add_heading_h2(doc, f"{num}. Conformités attestées")
    _add_body(doc, body or SEC_4_INTRO)


def _sec_4_1_docx(doc, body, parent_num=4, sub_num=1):
    _add_heading_h3(doc, f"{parent_num}.{sub_num} REACH — Substances "
                         f"extrêmement préoccupantes (SVHC)")
    _add_body(doc, body or SEC_4_1_BODY)


def _sec_4_2_docx(doc, body, parent_num=4, sub_num=2):
    _add_heading_h3(doc, f"{parent_num}.{sub_num} California Proposition 65")
    _add_body(doc, body or SEC_4_2_BODY)


def _sec_4_3_docx(doc, body, parent_num=4, sub_num=3):
    _add_heading_h3(doc, f"{parent_num}.{sub_num} Métaux lourds — traces "
                         f"techniquement inévitables")
    _add_body(doc, body or SEC_4_3_BODY)
    # Tableau métaux lourds
    tbl = doc.add_table(rows=5, cols=3)
    tbl.autofit = False
    headers = ["Élément", "Seuil individuel", "Seuil cumulé (94/62/CE)"]
    data_rows = [
        ("Plomb (Pb)", "< 50 ppm", "Somme Pb+Cd+Hg+Cr VI < 100 ppm"),
        ("Cadmium (Cd)", "< 25 ppm", ""),
        ("Mercure (Hg)", "< 12 ppm", ""),
        ("Chrome hexavalent (Cr VI)", "< 50 ppm", ""),
    ]
    # Header row
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        _set_cell_bg(c, LIGHT_BG_HEX)
        _set_cell_border(c)
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
    for i, row in enumerate(data_rows, start=1):
        for j, val in enumerate(row):
            c = tbl.rows[i].cells[j]
            _set_cell_border(c)
            c.text = ""
            p = c.paragraphs[0]
            if j == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = NAVY


def _sec_4_4_docx(doc, body, parent_num=4, sub_num=4):
    _add_heading_h3(doc, f"{parent_num}.{sub_num} Certification FSC")
    _add_body(doc, body or SEC_4_4_BODY)
    # Note : l'encadré jaune FSC est rendu par _sec_4_4_note_docx, section
    # pseudo « sec_4_4_note » (editable + removable indépendamment).


def _sec_4_4_note_docx(doc, body):
    """Encadré jaune « FSC en cours » — miroir DOCX de _sec_4_4_note (PDF)."""
    tbl = doc.add_table(rows=1, cols=1)
    c = tbl.rows[0].cells[0]
    _set_cell_bg(c, YELLOW_BG_HEX)
    _set_cell_border(c, color_hex=YELLOW_HEX, size="8")
    c.text = ""
    note_text = (body or SEC_4_4_NOTE)
    # Convertir HTML simplifié (SEC_4_4_NOTE contient <b>...</b>) en texte brut
    note_text = note_text.replace("<b>", "").replace("</b>", "")
    p = c.paragraphs[0]
    r = p.add_run(note_text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = NAVY


def _sec_4_5_docx(doc, body, parent_num=4, sub_num=5):
    _add_heading_h3(doc, f"{parent_num}.{sub_num} Absence de PFAS")
    _add_body(doc, body or SEC_4_5_BODY)


def _sec_4_6_docx(doc, body, parent_num=4, sub_num=6):
    _add_heading_h3(doc, f"{parent_num}.{sub_num} Absence de bisphénols "
                         f"(BPA, BPS, BPF)")
    _add_body(doc, body or SEC_4_6_BODY)


def _sec_4_7_docx(doc, body, parent_num=4, sub_num=7):
    _add_heading_h3(doc, f"{parent_num}.{sub_num} Certificats d'analyse "
                         f"laboratoire (CoA)")
    _add_body(doc, body or SEC_4_7_BODY)


def _sec_4_8_docx(doc, body, parent_num=4, sub_num=8):
    _add_heading_h3(doc, f"{parent_num}.{sub_num} Cadre général — PPWR")
    _add_body(doc, body or SEC_4_8_BODY)


def _sec_4_9_docx(doc, body, parent_num=4, sub_num=9):
    _add_heading_h3(doc, f"{parent_num}.{sub_num} Recyclabilité")
    _add_body(doc, body or SEC_4_9_BODY)


def _sec_5_docx(doc, body, num=5):
    _add_heading_h2(doc, f"{num}. Contenu recyclé")
    _add_body(doc, body or SEC_5_BODY)


def _sec_6_docx(doc, body, num=6):
    _add_heading_h2(doc, f"{num}. Base documentaire")
    _add_body(doc, body or SEC_6_BODY)


def _sec_7_docx(doc, body, ctx, num=7):
    _add_heading_h2(doc, f"{num}. Responsabilité et validité")
    validite_mois = int(ctx["validite_mois"] or 12)
    p = doc.add_paragraph()
    r = p.add_run(f"Validité : {validite_mois} mois à compter de la date "
                  f"d'émission.")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(4)
    _add_body(doc, body or SEC_7_BODY)


def _sec_8_docx(doc, body, ctx, num=8):
    _add_heading_h2(doc, f"{num}. Signature et cachet")
    _add_body(doc, body or SEC_8_BODY)

    # Priorité pour le Nom : version.representant → ctx.sig_nom → blanc.
    representant = (ctx.get("representant") or "").strip()
    sig_nom = representant or (ctx.get("sig_nom") or "").strip()
    sig_fonction = (ctx.get("sig_fonction") or "").strip()
    sig_date_raw = (ctx.get("sig_date") or "").strip()
    annee = ctx["annee"]

    # Date : vide → jour/mois/année d'émission ; sinon texte libre.
    date_line_txt = sig_date_raw or None
    if not date_line_txt:
        try:
            from datetime import datetime as _dt_sig
            d = _dt_sig.fromisoformat(ctx.get("date_emission_iso") or "")
            date_line_txt = f"{d.day:02d} / {d.month:02d} / {d.year}"
        except Exception:
            date_line_txt = f"____ / ____ / {annee}"

    # Tableau à 2 colonnes : signature à gauche, cachet à droite
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    left = tbl.rows[0].cells[0]
    right = tbl.rows[0].cells[1]
    left.width = Cm(9)
    right.width = Cm(8)
    left.text = ""
    right.text = ""

    # Colonne gauche : signature
    p0 = left.paragraphs[0]
    r0 = p0.add_run("Représentant SIFA")
    r0.bold = True
    r0.font.size = Pt(11)
    r0.font.color.rgb = NAVY

    def _add_kv_line(cell, label, value, bold_value):
        pp = cell.add_paragraph()
        rl = pp.add_run(label)
        rl.font.size = Pt(10)
        rl.font.color.rgb = NAVY
        rv = pp.add_run(value)
        rv.bold = bold_value
        rv.font.size = Pt(10)
        rv.font.color.rgb = NAVY

    if sig_nom:
        _add_kv_line(left, "Nom : ", sig_nom, True)
    else:
        _add_kv_line(left, "", "Nom : ______________________________", False)

    if sig_fonction:
        _add_kv_line(left, "Fonction : ", sig_fonction, True)
    else:
        _add_kv_line(left, "", "Fonction : __________________________", False)

    _add_kv_line(left, "Date : ", date_line_txt, bool(sig_date_raw))

    # Signature manuscrite : PNG si présent, sinon ligne vide pour signer
    p_sig_lbl = left.add_paragraph()
    r_sig_lbl = p_sig_lbl.add_run("Signature :")
    r_sig_lbl.font.size = Pt(10)
    r_sig_lbl.font.color.rgb = NAVY

    signature_path = _locate_signature()
    if signature_path:
        try:
            p_sig_img = left.add_paragraph()
            p_sig_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_s = p_sig_img.add_run()
            run_s.add_picture(signature_path, width=Mm(55))
        except Exception:
            pp = left.add_paragraph()
            rr = pp.add_run("______________________________________")
            rr.font.size = Pt(10)
            rr.font.color.rgb = NAVY
    else:
        left.add_paragraph()  # espace vertical
        pp = left.add_paragraph()
        rr = pp.add_run("______________________________________")
        rr.font.size = Pt(10)
        rr.font.color.rgb = NAVY

    # Colonne droite : cachet
    pr0 = right.paragraphs[0]
    rr0 = pr0.add_run("Cachet SIFA")
    rr0.bold = True
    rr0.font.size = Pt(11)
    rr0.font.color.rgb = NAVY

    cachet_path = _locate_cachet()
    if cachet_path:
        try:
            pcachet = right.add_paragraph()
            pcachet.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_c = pcachet.add_run()
            run_c.add_picture(cachet_path, width=Mm(55))
        except Exception:
            pfallback = right.add_paragraph()
            rf = pfallback.add_run("[Emplacement réservé au cachet]")
            rf.italic = True
            rf.font.size = Pt(9.5)
            rf.font.color.rgb = MUTED
    else:
        pfallback = right.add_paragraph()
        rf = pfallback.add_run("[Emplacement réservé au cachet]")
        rf.italic = True
        rf.font.size = Pt(9.5)
        rf.font.color.rgb = MUTED

    # Footer typo
    doc.add_paragraph()
    pfoot = doc.add_paragraph()
    pfoot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = pfoot.add_run("Document établi conformément à l'Annexe VIII du "
                       "règlement (UE) 2025/40 (PPWR).")
    r1.italic = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = MUTED
    pfoot2 = doc.add_paragraph()
    pfoot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = pfoot2.add_run(f"Référence unique : {ctx.get('ref') or ''} — "
                        f"Version 1.0")
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = MUTED


# ─── Orchestrateur ────────────────────────────────────────────────────────
def _build_docx(ctx, sections_overrides: dict,
                template_default_overrides: dict) -> bytes:
    doc = Document()
    # Style de base
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = NAVY

    _install_header_footer(doc, ref=ctx["ref"], header_date=ctx["header_date"])
    _title_block_docx(doc, ref=ctx["ref"],
                      date_emission_iso=ctx["date_emission_iso"])

    ov_all = sections_overrides or {}
    tpl_ov_all = template_default_overrides or {}

    # Pré-résolution des champs de signature (sig_nom, sig_fonction, sig_date).
    # Même priorité que les body overrides : version > template > default_body.
    for sec in SECTIONS_META:
        if not sec.get("is_signature_field"):
            continue
        sid = sec["id"]
        resolved = ""
        v_ov = ov_all.get(sid, {})
        v_custom = v_ov.get("custom_body") if isinstance(v_ov, dict) else None
        if v_custom and str(v_custom).strip():
            resolved = str(v_custom).strip()
        else:
            t_default = tpl_ov_all.get(sid)
            if t_default and str(t_default).strip():
                resolved = str(t_default).strip()
            else:
                resolved = sec.get("default_body") or ""
        ctx[sid] = resolved

    # Renumérotation dynamique — miroir de sifa_doc_pdf._build_flowables
    main_counter = 0
    main_num_by_id = {}
    for sec in SECTIONS_META:
        if not sec.get("is_main"):
            continue
        if not _section_included(sec, ov_all):
            continue
        main_counter += 1
        main_num_by_id[sec["id"]] = main_counter

    sub_counters = {}
    sub_info_by_id = {}
    for sec in SECTIONS_META:
        parent_id = sec.get("parent")
        if not parent_id:
            continue
        if not _section_included(sec, ov_all):
            continue
        # Notes inline (encadrés) : rendues sous leur parent sans numéro propre.
        if sec.get("is_note"):
            continue
        # Champs signature : pré-résolus dans ctx, non rendus comme sections.
        if sec.get("is_signature_field"):
            continue
        parent_num = main_num_by_id.get(parent_id)
        if parent_num is None:
            continue
        sub_counters[parent_id] = sub_counters.get(parent_id, 0) + 1
        sub_info_by_id[sec["id"]] = {
            "parent_num": parent_num,
            "sub_num": sub_counters[parent_id],
        }

    # Rendu
    for sec in SECTIONS_META:
        if not _section_included(sec, ov_all):
            continue
        # Champs signature : pas de rendu, injectés dans _sec_8_docx via ctx.
        if sec.get("is_signature_field"):
            continue
        parent_id = sec.get("parent")
        if parent_id and parent_id not in main_num_by_id:
            continue

        # Résolution du body (uniquement pour les sections editable)
        default_body = sec.get("default_body") or ""
        body = None
        if sec.get("editable"):
            body = _resolve_body(sec["id"], default_body, ov_all, tpl_ov_all)

        sid = sec["id"]
        if sec.get("is_main"):
            num = main_num_by_id[sid]
            if sid == "sec_1":
                _sec_1_docx(doc, num=num)
            elif sid == "sec_2":
                _sec_2_docx(doc, body, num=num)
            elif sid == "sec_3":
                _sec_3_docx(doc, ctx, num=num)
            elif sid == "sec_4_intro":
                _sec_4_intro_docx(doc, body, num=num)
            elif sid == "sec_5":
                _sec_5_docx(doc, body, num=num)
            elif sid == "sec_6":
                _sec_6_docx(doc, body, num=num)
            elif sid == "sec_7":
                _sec_7_docx(doc, body, ctx, num=num)
            elif sid == "sec_8":
                _sec_8_docx(doc, body, ctx, num=num)
        elif sid in sub_info_by_id:
            info = sub_info_by_id[sid]
            args = (doc, body)
            kw = {"parent_num": info["parent_num"], "sub_num": info["sub_num"]}
            if sid == "sec_4_1":
                _sec_4_1_docx(*args, **kw)
            elif sid == "sec_4_2":
                _sec_4_2_docx(*args, **kw)
            elif sid == "sec_4_3":
                _sec_4_3_docx(*args, **kw)
            elif sid == "sec_4_4":
                _sec_4_4_docx(*args, **kw)
            elif sid == "sec_4_5":
                _sec_4_5_docx(*args, **kw)
            elif sid == "sec_4_6":
                _sec_4_6_docx(*args, **kw)
            elif sid == "sec_4_7":
                _sec_4_7_docx(*args, **kw)
            elif sid == "sec_4_8":
                _sec_4_8_docx(*args, **kw)
            elif sid == "sec_4_9":
                _sec_4_9_docx(*args, **kw)
        elif sec.get("is_note"):
            # Notes inline (encadrés) — rendues sans numéro propre, juste
            # après leur sous-section parente dans l'ordre de SECTIONS_META.
            if sid == "sec_4_4_note":
                _sec_4_4_note_docx(doc, body)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── API publique ─────────────────────────────────────────────────────────
def build_declaration_ue_docx(*, client_nom: str, fournisseurs: list,
                              ref: str, date_emission_iso: str,
                              validite_mois: int = 12,
                              sections_overrides: dict = None,
                              template_default_overrides: dict = None,
                              representant: str = None) -> bytes:
    """
    Génère un DOCX de Déclaration UE de Conformité — miroir éditable du PDF.

    Signature identique à build_declaration_ue_pdf.

    Returns:
        bytes: contenu binaire du fichier .docx
    """
    header_date = _fr_month_year(date_emission_iso)
    _, _, annee = _fr_date_parts(date_emission_iso)
    ctx = {
        "client_nom": client_nom or "",
        "fournisseurs": fournisseurs or [],
        "ref": ref,
        "date_emission_iso": date_emission_iso,
        "validite_mois": validite_mois,
        "annee": annee,
        "header_date": header_date,
        "representant": representant or "",
    }
    return _build_docx(ctx, sections_overrides or {},
                       template_default_overrides or {})


# ─── Registry générique multi-templates ──────────────────────────────────
TEMPLATE_BUILDERS_DOCX = {
    "declaration_ue": build_declaration_ue_docx,
}


def build_template_docx(template_code: str, **kwargs) -> bytes:
    """Point d'entrée générique pour construire un DOCX à partir du code
    template. Symétrique à sifa_doc_pdf.build_template_pdf."""
    builder = TEMPLATE_BUILDERS_DOCX.get(template_code)
    if not builder:
        raise ValueError(f"Template DOCX inconnu: {template_code}")
    return builder(**kwargs)
